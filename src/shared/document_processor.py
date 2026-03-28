"""
Document Processor - Parse and chunk documents for knowledge ingestion.

Supports: PDF, DOCX, TXT, MD, HTML, XLSX, XLS, CSV, TSV files
Chunking strategy: Semantic (section/header-based) with smart section detection
Spreadsheets: Converted to row-based text chunks for semantic search
"""

import os
import re
import uuid
from typing import Optional, List
from dataclasses import dataclass, field


@dataclass
class DocumentChunk:
    """A single chunk of document content."""
    text: str
    chunk_index: int
    chunk_type: str  # section, paragraph, table, contact_info
    section_title: Optional[str] = None
    page_number: Optional[int] = None
    metadata: dict = field(default_factory=dict)


# Common section headers in resumes and documents
RESUME_SECTIONS = [
    r'^(education|educational\s*background|academics?)\s*$',
    r'^(experience|work\s*experience|employment|professional\s*experience|work\s*history)\s*$',
    r'^(skills|technical\s*skills|competencies|expertise|tech\s*stack)\s*$',
    r'^(projects|personal\s*projects|key\s*projects)\s*$',
    r'^(certifications?|licenses?)\s*$',
    r'^(publications?|research)\s*$',
    r'^(awards?|honors?|achievements)\s*$',
    r'^(languages?)\s*$',
    r'^(interests?|hobbies?|activities)\s*$',
    r'^(summary|profile|objective|about\s*me)\s*$',
    r'^(contact|personal\s*info|personal\s*details)\s*$',
]


class DocumentProcessor:
    """Parse and chunk documents for knowledge ingestion."""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.md', '.html', '.xlsx', '.xls', '.csv', '.tsv'}
    SPREADSHEET_EXTENSIONS = {'.xlsx', '.xls', '.csv', '.tsv'}
    
    def __init__(self):
        self.max_chunk_tokens = 1500
        self.min_chunk_tokens = 50
    
    def get_file_type(self, file_path: str) -> Optional[str]:
        """Detect file type from extension."""
        ext = os.path.splitext(file_path)[1].lower()
        if ext in self.SUPPORTED_EXTENSIONS:
            return ext.lstrip('.')
        return None
    
    def extract_text(self, file_path: str) -> tuple[str, Optional[str]]:
        """
        Extract text from file.
        Returns (text, error_message).
        """
        file_type = self.get_file_type(file_path)
        if not file_type:
            return "", f"Unsupported file type: {file_path}"
        
        try:
            if file_type == 'pdf':
                return self._extract_pdf(file_path), None
            elif file_type in ('docx', 'doc'):
                return self._extract_docx(file_path), None
            elif file_type in ('txt', 'md'):
                return self._extract_text_file(file_path), None
            elif file_type == 'html':
                return self._extract_html(file_path), None
            elif file_type in ('xlsx', 'xls', 'csv', 'tsv'):
                return self._extract_spreadsheet(file_path), None
        except Exception as e:
            return "", f"Error extracting {file_path}: {e}"
        
        return "", "Unknown file type"
    
    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF using pypdf with page markers."""
        from pypdf import PdfReader
        
        reader = PdfReader(file_path)
        text_parts = []
        
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text() or ""
            if page_text.strip():
                text_parts.append(f"[PAGE:{i+1}]\n{page_text}")
        
        return "\n\n".join(text_parts)
    
    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX using python-docx."""
        from docx import Document
        
        doc = Document(file_path)
        text_parts = []
        
        for para in doc.paragraphs:
            if para.text.strip():
                if para.style.name.startswith('Heading'):
                    level = para.style.name.replace('Heading ', '')
                    prefix = '#' * int(level) if level.isdigit() else '#'
                    text_parts.append(f"\n{prefix} {para.text}\n")
                else:
                    text_parts.append(para.text)
        
        for table in doc.tables:
            table_text = []
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                table_text.append(row_text)
            if table_text:
                text_parts.append("\n[TABLE]\n" + "\n".join(table_text) + "\n[/TABLE]\n")
        
        return "\n".join(text_parts)
    
    def _extract_text_file(self, file_path: str) -> str:
        """Read plain text or markdown file."""
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            return f.read()
    
    def _extract_html(self, file_path: str) -> str:
        """Extract text from HTML using BeautifulSoup."""
        from bs4 import BeautifulSoup
        
        with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
            soup = BeautifulSoup(f.read(), 'html.parser')
        
        for element in soup(['script', 'style', 'nav', 'footer']):
            element.decompose()
        
        return soup.get_text(separator='\n', strip=True)
    
    def _extract_spreadsheet(self, file_path: str) -> str:
        """Extract text from spreadsheet files (XLSX, XLS, CSV, TSV)."""
        from src.shared.spreadsheet_processor import SpreadsheetProcessor
        
        processor = SpreadsheetProcessor()
        info, rows = processor.parse_file(file_path)
        
        text_parts = []
        text_parts.append(f"[SPREADSHEET: {info.name}]")
        text_parts.append(f"Type: {info.file_type}")
        if info.sheet_name:
            text_parts.append(f"Sheet: {info.sheet_name}")
        text_parts.append(f"Rows: {info.row_count}, Columns: {info.column_count}")
        text_parts.append(f"Headers: {', '.join(info.column_headers)}")
        text_parts.append("")
        
        for row in rows:
            row_text = processor.generate_row_text(row, info.column_headers)
            text_parts.append(f"[ROW:{row.row_index}] {row_text}")
        
        return "\n".join(text_parts)
    
    def _detect_section_header(self, line: str) -> Optional[str]:
        """Detect if a line is a section header."""
        line_lower = line.strip().lower()
        
        # Check for markdown headers
        if line.startswith('#'):
            return line.lstrip('#').strip()
        
        # Check for common resume section patterns
        for pattern in RESUME_SECTIONS:
            if re.match(pattern, line_lower, re.IGNORECASE):
                return line.strip()
        
        # Check for ALL CAPS headers (common in resumes)
        if line.isupper() and len(line.strip()) < 50:
            return line.strip()
        
        # Check for Title Case headers followed by colon or standalone
        words = line.strip().split()
        if 2 <= len(words) <= 5:
            if all(w[0].isupper() or w.lower() in ['and', 'or', 'of', 'for', 'the', 'in', 'to'] for w in words if w):
                return line.strip()
        
        return None
    
    def chunk_text(self, text: str, file_type: str = None) -> List[DocumentChunk]:
        """
        Split text into semantic chunks with smart section detection.
        
        Strategy:
        1. Detect section headers (markdown, resume patterns, ALL CAPS)
        2. Split by detected sections
        3. For large sections, apply paragraph-based splitting
        """
        chunks = []
        lines = text.split('\n')
        
        # First pass: detect sections
        sections = []
        current_section = {"title": None, "lines": [], "start_page": None}
        
        for line in lines:
            # Check for page markers
            page_match = re.match(r'\[PAGE:(\d+)\]', line)
            if page_match:
                current_section["start_page"] = int(page_match.group(1))
                continue
            
            # Check for section header
            header = self._detect_section_header(line)
            
            if header and len(line.strip()) < 100:  # Headers are typically short
                # Save current section if it has content
                if current_section["lines"]:
                    sections.append(current_section.copy())
                current_section = {"title": header, "lines": [], "start_page": current_section.get("start_page")}
            else:
                current_section["lines"].append(line)
        
        # Don't forget the last section
        if current_section["lines"]:
            sections.append(current_section)
        
        # If no sections detected, try markdown-style headers
        if len(sections) <= 1:
            sections = self._chunk_by_markdown_headers(text)
        
        # If still no sections, fall back to paragraph-based
        if len(sections) <= 1:
            return self._chunk_by_paragraphs(text)
        
        # Second pass: create chunks from sections
        chunk_index = 0
        for section in sections:
            section_text = '\n'.join(section["lines"]).strip()
            
            if not section_text:
                continue
            
            # Check if section is too large
            approx_tokens = len(section_text) // 4
            
            if approx_tokens > self.max_chunk_tokens:
                # Split large sections by paragraphs
                sub_chunks = self._split_large_section(
                    section_text, 
                    section["title"], 
                    chunk_index,
                    section.get("start_page")
                )
                chunks.extend(sub_chunks)
                chunk_index += len(sub_chunks)
            elif len(section_text) >= self.min_chunk_tokens:
                chunks.append(DocumentChunk(
                    text=section_text,
                    chunk_index=chunk_index,
                    chunk_type='section',
                    section_title=section["title"],
                    page_number=section.get("start_page")
                ))
                chunk_index += 1
        
        return chunks if chunks else self._chunk_by_paragraphs(text)
    
    def _chunk_by_markdown_headers(self, text: str) -> List[dict]:
        """Split by markdown-style headers."""
        sections = []
        header_pattern = r'\n(#{1,3})\s+(.+?)(?=\n)'
        headers = list(re.finditer(header_pattern, text))
        
        if not headers:
            return [{"title": None, "lines": text.split('\n'), "start_page": None}]
        
        for i, match in enumerate(headers):
            start = match.end()
            end = headers[i + 1].start() if i + 1 < len(headers) else len(text)
            section_text = text[start:end].strip()
            
            sections.append({
                "title": match.group(2).strip(),
                "lines": section_text.split('\n'),
                "start_page": None
            })
        
        return sections
    
    def _chunk_by_paragraphs(self, text: str) -> List[DocumentChunk]:
        """Fallback: Split by paragraphs with max token limit."""
        paragraphs = re.split(r'\n\s*\n', text)
        chunks = []
        current_chunk = ""
        chunk_index = 0
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            para_tokens = len(para) // 4
            
            if len(current_chunk) // 4 + para_tokens > self.max_chunk_tokens:
                if current_chunk:
                    chunks.append(DocumentChunk(
                        text=current_chunk.strip(),
                        chunk_index=chunk_index,
                        chunk_type='paragraph'
                    ))
                    chunk_index += 1
                    current_chunk = ""
            
            current_chunk += para + "\n\n"
        
        if current_chunk.strip():
            chunks.append(DocumentChunk(
                text=current_chunk.strip(),
                chunk_index=chunk_index,
                chunk_type='paragraph'
            ))
        
        return chunks
    
    def _split_large_section(
        self, 
        text: str, 
        section_title: str, 
        start_index: int,
        page_number: Optional[int] = None
    ) -> List[DocumentChunk]:
        """Split a large section into smaller chunks."""
        chunks = []
        paragraphs = re.split(r'\n\s*\n', text)
        current_chunk = ""
        chunk_index = start_index
        
        for para in paragraphs:
            para = para.strip()
            if not para:
                continue
            
            if len(current_chunk) // 4 + len(para) // 4 > self.max_chunk_tokens:
                if current_chunk:
                    chunks.append(DocumentChunk(
                        text=current_chunk.strip(),
                        chunk_index=chunk_index,
                        chunk_type='section',
                        section_title=section_title,
                        page_number=page_number
                    ))
                    chunk_index += 1
                    current_chunk = ""
            
            current_chunk += para + "\n\n"
        
        if current_chunk.strip():
            chunks.append(DocumentChunk(
                text=current_chunk.strip(),
                chunk_index=chunk_index,
                chunk_type='section',
                section_title=section_title,
                page_number=page_number
            ))
        
        return chunks
    
    def process_file(
        self,
        file_path: str,
        doc_name: Optional[str] = None,
        category: Optional[str] = None,
        tags: list = None,
        source: str = "cli"
    ) -> tuple[str, list]:
        """
        Full pipeline: extract → chunk → return doc_id and chunks.
        
        Returns (doc_id, list of chunk dicts ready for storage).
        """
        doc_id = str(uuid.uuid4())[:8]
        doc_name = doc_name or os.path.basename(file_path)
        file_type = self.get_file_type(file_path) or "unknown"
        
        text, error = self.extract_text(file_path)
        if error:
            raise ValueError(error)
        
        chunks = self.chunk_text(text, file_type)
        
        chunk_dicts = []
        for chunk in chunks:
            chunk_dict = {
                "text": chunk.text,
                "chunk_index": chunk.chunk_index,
                "chunk_type": chunk.chunk_type,
                "section_title": chunk.section_title,
                "doc_id": doc_id,
                "doc_name": doc_name,
                "file_type": file_type,
                "category": category,
                "source": source,
                "tags": tags or [],
            }
            chunk_dicts.append(chunk_dict)
        
        return doc_id, chunk_dicts


def process_document(
    file_path: str,
    doc_name: Optional[str] = None,
    category: Optional[str] = None,
    tags: list = None,
    source: str = "cli"
) -> tuple[str, list]:
    """Convenience function to process a document."""
    processor = DocumentProcessor()
    return processor.process_file(file_path, doc_name, category, tags, source)
